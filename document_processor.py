"""
Complete Medical Document & Image Processor
Supports: PDF, DOCX, TXT, JSON, Images (JPG, PNG, etc.) with OCR
"""

import os
import json
import sqlite3
import re
from datetime import datetime
from pathlib import Path
from PIL import Image
import pytesseract
import cv2
import numpy as np
import PyPDF2
import docx


pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class MedicalDocumentProcessor:
    """Unified processor for all medical documents and images"""
    
    def __init__(self, db_path='medical_records.db'):
        self.db_path = db_path
        self.storage_path = 'medical_documents'
        
        # File format support
        self.supported_text_formats = ['.txt', '.pdf', '.docx', '.json']
        self.supported_image_formats = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
        
        # Performance settings
        self.max_pdf_pages = 100
        self.max_text_length = 1_000_000
        self.max_image_dimension = 2048
        
        self.init_database()
        self.ensure_storage_directory()
    
    def init_database(self):
        """Initialize database with comprehensive schema"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS patients (
                patient_id TEXT PRIMARY KEY,
                name TEXT,
                date_of_birth TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS medical_documents (
                document_id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_id TEXT,
                file_path TEXT,
                file_type TEXT,
                file_format TEXT,
                upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                extracted_text TEXT,
                metadata TEXT,
                report_type TEXT,
                doctor_name TEXT,
                report_date TEXT,
                department TEXT,
                diagnosis TEXT,
                medications TEXT,
                FOREIGN KEY (patient_id) REFERENCES patients(patient_id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS medical_records (
                record_id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_id TEXT,
                record_type TEXT,
                content TEXT,
                doctor_name TEXT,
                record_date TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                document_id INTEGER,
                FOREIGN KEY (patient_id) REFERENCES patients(patient_id),
                FOREIGN KEY (document_id) REFERENCES medical_documents(document_id)
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def ensure_storage_directory(self):
        """Create storage directory if it doesn't exist"""
        if not os.path.exists(self.storage_path):
            os.makedirs(self.storage_path)
    
    # ==================== TEXT DOCUMENT PROCESSORS ====================
    
    def process_txt(self, file_path):
        """Process plain text file"""
        file_path = Path(file_path)
        file_size = file_path.stat().st_size
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            if file_size > self.max_text_length:
                text = f.read(self.max_text_length)
                truncated = True
            else:
                text = f.read()
                truncated = False
        
        return {
            'type': 'text',
            'text': text,
            'metadata': {
                'filename': file_path.name,
                'format': 'txt',
                'size_bytes': file_size,
                'truncated': truncated
            }
        }
    
    def process_pdf(self, file_path):
        """Process PDF file with page limit"""
        file_path = Path(file_path)
        text_content = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                pages_to_process = min(num_pages, self.max_pdf_pages)
                truncated = num_pages > self.max_pdf_pages
                
                for page_num in range(pages_to_process):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(text)
                        
                        current_length = sum(len(t) for t in text_content)
                        if current_length > self.max_text_length:
                            truncated = True
                            break
                            
                    except Exception as e:
                        print(f"⚠️ Warning: Could not extract page {page_num}: {e}")
                        continue
        
        except Exception as e:
            raise ValueError(f"Error processing PDF: {e}")
        
        if not text_content:
            raise ValueError("No text could be extracted from PDF")
        
        full_text = '\n\n'.join(text_content)
        
        return {
            'type': 'text',
            'text': full_text[:self.max_text_length],
            'metadata': {
                'filename': file_path.name,
                'format': 'pdf',
                'num_pages': num_pages,
                'pages_processed': pages_to_process,
                'size_bytes': file_path.stat().st_size,
                'truncated': truncated
            }
        }
    
    def process_docx(self, file_path):
        """Process Word document"""
        file_path = Path(file_path)
        
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            total_chars = 0
            truncated = False
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    total_chars += len(paragraph.text)
                    
                    if total_chars > self.max_text_length:
                        truncated = True
                        break
            
            if not text_content:
                raise ValueError("No text content found in document")
            
            return {
                'type': 'text',
                'text': '\n\n'.join(text_content),
                'metadata': {
                    'filename': file_path.name,
                    'format': 'docx',
                    'num_paragraphs': len(doc.paragraphs),
                    'size_bytes': file_path.stat().st_size,
                    'truncated': truncated
                }
            }
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {e}")
    
    def process_json(self, file_path):
        """Process JSON file"""
        file_path = Path(file_path)
        file_size = file_path.stat().st_size
        
        if file_size > 10 * 1024 * 1024:
            raise ValueError(f"JSON file too large: {file_size / (1024*1024):.1f}MB")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file: {e}")
        
        if isinstance(data, dict) and 'text' in data:
            text = data['text']
            metadata = {k: v for k, v in data.items() if k != 'text'}
        else:
            text = json.dumps(data, indent=2)
            metadata = {}
        
        if len(text) > self.max_text_length:
            text = text[:self.max_text_length]
            metadata['truncated'] = True
        
        metadata.update({
            'filename': file_path.name,
            'format': 'json',
            'size_bytes': file_size
        })
        
        return {
            'type': 'text',
            'text': text,
            'metadata': metadata
        }
    
    # ==================== IMAGE PROCESSORS WITH OCR ====================
    
    def preprocess_image_for_ocr(self, image_path):
        """Preprocess image for better OCR results"""
        img = cv2.imread(str(image_path))
        
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.fastNlMeansDenoising(gray)
        thresh = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY, 11, 2
        )
        
        temp_path = str(image_path).replace('.', '_processed.')
        cv2.imwrite(temp_path, thresh)
        
        return temp_path
    
    def extract_text_from_image(self, image_path):
        """Extract text from image using OCR"""
        try:
            processed_path = self.preprocess_image_for_ocr(image_path)
            
            text = pytesseract.image_to_string(
                Image.open(processed_path),
                lang='eng',
                config='--psm 6'
            )
            
            if os.path.exists(processed_path):
                os.remove(processed_path)
            
            return text.strip()
        except Exception as e:
            print(f"⚠️ OCR Error: {e}")
            return ""
    
    def process_image(self, file_path):
        """Process image file with OCR"""
        file_path = Path(file_path)
        
        try:
            with Image.open(file_path) as img:
                original_format = img.format
                original_size = img.size
                original_mode = img.mode
                
                width, height = img.size
                needs_resize = width > self.max_image_dimension or height > self.max_image_dimension
                
                if needs_resize:
                    if width > height:
                        new_width = self.max_image_dimension
                        new_height = int(height * (self.max_image_dimension / width))
                    else:
                        new_height = self.max_image_dimension
                        new_width = int(width * (self.max_image_dimension / height))
                    resized = (new_width, new_height)
                else:
                    resized = None
            
            # Extract text using OCR
            extracted_text = self.extract_text_from_image(file_path)
            
            return {
                'type': 'image',
                'text': extracted_text,
                'path': str(file_path.absolute()),
                'metadata': {
                    'filename': file_path.name,
                    'format': original_format,
                    'size': original_size,
                    'mode': original_mode,
                    'size_bytes': file_path.stat().st_size,
                    'needs_resize': needs_resize,
                    'suggested_size': resized,
                    'text_extracted': len(extracted_text) > 0
                }
            }
        except Exception as e:
            raise ValueError(f"Invalid image file: {e}")
    
    # ==================== METADATA EXTRACTION ====================
    
    def extract_medical_metadata(self, text, fast_mode=True):
        """Extract medical metadata from text"""
        if fast_mode or len(text) > 5000:
            text_to_analyze = text[:5000]
        else:
            text_to_analyze = text
        
        metadata = {
            'report_type': 'general',
            'doctor': None,
            'diagnosis': None,
            'medications': [],
            'department': None,
            'report_date': None
        }
        
        text_lower = text_to_analyze.lower()
        
        # Detect report type
        report_keywords = {
            'lab_results': ['lab', 'laboratory', 'test result'],
            'imaging': ['x-ray', 'mri', 'ct scan', 'ultrasound', 'imaging'],
            'diagnosis': ['diagnosis', 'diagnosed with'],
            'follow_up': ['follow-up', 'followup', 'follow up'],
            'emergency': ['emergency', 'er', 'acute', 'urgent'],
            'prescription': ['prescription', 'rx', 'prescribed'],
            'consultation': ['consultation', 'visit', 'checkup', 'check-up']
        }
        
        for report_type, keywords in report_keywords.items():
            if any(word in text_lower for word in keywords):
                metadata['report_type'] = report_type
                break
        
        # Extract doctor name
        doctor_patterns = [
            r'(?:Dr\.?|Doctor)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'Physician:\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'Consultant:\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
        ]
        for pattern in doctor_patterns:
            match = re.search(pattern, text_to_analyze, re.IGNORECASE)
            if match:
                metadata['doctor'] = match.group(1).strip()
                break
        
        # Extract dates
        date_patterns = [
            r'(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})',
            r'Date:\s*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})'
        ]
        for pattern in date_patterns:
            match = re.search(pattern, text_to_analyze, re.IGNORECASE)
            if match:
                metadata['report_date'] = match.group(1).strip()
                break
        
        # Extract department
        dept_keywords = ['Cardiology', 'Radiology', 'Pathology', 'Neurology', 
                        'Orthopedics', 'Surgery', 'Medicine', 'Emergency']
        for dept in dept_keywords:
            if dept.lower() in text_lower:
                metadata['department'] = dept
                break
        
        # Extract diagnosis
        diagnosis_patterns = [
            r'diagnosis:?\s*([^\n]{1,200})',
            r'diagnosed\s+with:?\s*([^\n]{1,200})',
            r'impression:?\s*([^\n]{1,200})',
        ]
        for pattern in diagnosis_patterns:
            match = re.search(pattern, text_lower)
            if match:
                diagnosis = match.group(1).strip()
                diagnosis = ' '.join(diagnosis.split())
                metadata['diagnosis'] = diagnosis[:200]
                break
        
        # Extract medications
        med_text_section = text_lower[:3000]
        med_patterns = [
            r'medications?:?\s*([^\n]+(?:\n[^\n]+){0,5})',
            r'prescribed:?\s*([^\n]+(?:\n[^\n]+){0,5})',
            r'drugs?:?\s*([^\n]+(?:\n[^\n]+){0,5})',
        ]
        for pattern in med_patterns:
            match = re.search(pattern, med_text_section, re.MULTILINE)
            if match:
                meds_text = match.group(1)
                medications = re.split(r'[,;\n•\-]', meds_text)
                cleaned_meds = []
                for med in medications:
                    med = med.strip()
                    if med and len(med) > 2 and len(med) < 50:
                        if not any(skip in med.lower() for skip in ['none', 'n/a', 'patient', 'see']):
                            cleaned_meds.append(med.title())
                    if len(cleaned_meds) >= 10:
                        break
                
                metadata['medications'] = cleaned_meds
                break
        
        return metadata
    
    # ==================== UNIFIED UPLOAD SYSTEM ====================
    
    def upload_document(self, patient_id, file_path):
        """Upload any supported document type"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"success": False, "error": "File not found"}
            
            # Check file size
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > 100:
                return {"success": False, "error": f"File too large: {file_size_mb:.1f}MB (max 100MB)"}
            
            extension = file_path.suffix.lower()
            
            # Process based on file type
            print(f"📄 Processing {extension} file...")
            
            if extension == '.txt':
                result = self.process_txt(file_path)
            elif extension == '.pdf':
                result = self.process_pdf(file_path)
            elif extension == '.docx':
                result = self.process_docx(file_path)
            elif extension == '.json':
                result = self.process_json(file_path)
            elif extension in self.supported_image_formats:
                result = self.process_image(file_path)
            else:
                return {"success": False, "error": f"Unsupported file format: {extension}"}
            
            # Copy file to storage
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{patient_id}_{timestamp}_{file_path.name}"
            storage_path = os.path.join(self.storage_path, filename)
            
            import shutil
            shutil.copy2(file_path, storage_path)
            
            # Extract metadata
            print(f"🔍 Extracting medical metadata...")
            extracted_text = result.get('text', '')
            medical_metadata = self.extract_medical_metadata(extracted_text)
            
            # Save to database
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO medical_documents 
                (patient_id, file_path, file_type, file_format, extracted_text, 
                 metadata, report_type, doctor_name, report_date, department, 
                 diagnosis, medications)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                patient_id,
                storage_path,
                result['type'],
                result['metadata']['format'],
                extracted_text,
                json.dumps(result['metadata']),
                medical_metadata['report_type'],
                medical_metadata['doctor'],
                medical_metadata['report_date'],
                medical_metadata['department'],
                medical_metadata['diagnosis'],
                json.dumps(medical_metadata['medications'])
            ))
            
            document_id = cursor.lastrowid
            
            # Create medical record entry
            cursor.execute('''
                INSERT INTO medical_records
                (patient_id, record_type, content, doctor_name, record_date, document_id)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                patient_id,
                f'{result["type"].title()} - {medical_metadata["report_type"]}',
                extracted_text[:500] + '...' if len(extracted_text) > 500 else extracted_text,
                medical_metadata['doctor'],
                medical_metadata['report_date'],
                document_id
            ))
            
            conn.commit()
            conn.close()
            
            return {
                "success": True,
                "document_id": document_id,
                "file_type": result['type'],
                "format": result['metadata']['format'],
                "extracted_text_length": len(extracted_text),
                "metadata": medical_metadata,
                "storage_path": storage_path
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def batch_upload_folder(self, patient_id, folder_path):
        """Upload all supported files from a folder"""
        folder_path = Path(folder_path)
        
        if not folder_path.exists():
            return {"success": False, "error": "Folder not found"}
        
        all_formats = self.supported_text_formats + self.supported_image_formats
        files = [f for f in folder_path.iterdir() 
                if f.suffix.lower() in all_formats]
        
        print(f"📁 Found {len(files)} files to process...")
        
        results = []
        for idx, file_path in enumerate(files, 1):
            print(f"\n[{idx}/{len(files)}] Processing {file_path.name}...")
            result = self.upload_document(patient_id, file_path)
            results.append({
                'filename': file_path.name,
                'result': result
            })
        
        successful = sum(1 for r in results if r['result'].get('success'))
        
        return {
            "success": True,
            "total_files": len(files),
            "successful": successful,
            "failed": len(files) - successful,
            "results": results
        }
    
    def search_patient_records(self, patient_id, search_query=None, months_filter=None):
        """Search patient records with filters"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        query = '''
            SELECT r.record_id, r.record_type, r.content, r.doctor_name, 
                   r.record_date, r.created_at, d.file_path, d.extracted_text
            FROM medical_records r
            LEFT JOIN medical_documents d ON r.document_id = d.document_id
            WHERE r.patient_id = ?
        '''
        params = [patient_id]
        
        if months_filter:
            query += ' AND r.created_at >= datetime("now", ?)'
            params.append(f'-{months_filter} months')
        
        if search_query:
            query += ''' AND (
                LOWER(r.content) LIKE ? OR 
                LOWER(r.doctor_name) LIKE ? OR
                LOWER(d.extracted_text) LIKE ?
            )'''
            search_param = f'%{search_query.lower()}%'
            params.extend([search_param, search_param, search_param])
        
        query += ' ORDER BY r.created_at DESC'
        
        cursor.execute(query, params)
        results = cursor.fetchall()
        conn.close()
        
        return results


# Example usage
if __name__ == "__main__":
    processor = MedicalDocumentProcessor()
    
    print("=" * 70)
    print("📋 UNIFIED MEDICAL DOCUMENT PROCESSOR")
    print("=" * 70)
    print("\nSupported formats:")
    print("  📄 Text: TXT, PDF, DOCX, JSON")
    print("  🖼️  Images: JPG, PNG, BMP, TIFF (with OCR)")
    
    patient_id = input("\nEnter Patient ID: ").strip()
    file_path = input("Enter file/folder path: ").strip()
    
    if os.path.isdir(file_path):
        # Batch upload
        print("\n📁 Batch upload mode...")
        result = processor.batch_upload_folder(patient_id, file_path)
        print(f"\n✅ Processed {result['successful']}/{result['total_files']} files successfully")
    
    elif os.path.isfile(file_path):
        # Single file upload
        result = processor.upload_document(patient_id, file_path)
        
        if result['success']:
            print("\n✅ Document uploaded successfully!")
            print(f"📄 Type: {result['file_type']} ({result['format']})")
            print(f"📝 Extracted: {result['extracted_text_length']} characters")
            print(f"📋 Report Type: {result['metadata']['report_type']}")
            print(f"👨‍⚕️ Doctor Name: {result['metadata']['doctor'] or 'Not found'}")
            print(f"📅 Date: {result['metadata']['report_date'] or 'Not found'}")
            print(f"🏥 Medicines: {result['metadata']['department'] or 'Not found'}")
        else:
            print(f"\n❌ Upload failed: {result['error']}")
    else:
        print("❌ File or folder not found!")